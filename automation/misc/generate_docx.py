import argparse
import os
import random
from pathlib import Path

from docx import Document
from file_utils import system_path


def generate_meaningful_or_tech_text(paragraphs=3):
	"""
	Generate either meaningful or tech-related English text for the documents.
	"""
	motivational_texts = [
		"Success is not the key to happiness. Happiness is the key to success. If you love what you are doing, you will be successful.",
		"Life is not measured by the number of breaths we take, but by the moments that take our breath away.",
		"The greatest glory in living lies not in never falling, but in rising every time we fall. - Nelson Mandela",
		"Do not dwell in the past, do not dream of the future, concentrate the mind on the present moment. - Buddha",
		"Believe you can and you're halfway there. - Theodore Roosevelt",
		"In the middle of every difficulty lies opportunity. - Albert Einstein",
		"The journey of a thousand miles begins with a single step. - Lao Tzu",
	]
	
	tech_texts = [
		"The difference between synchronous and asynchronous programming lies in how tasks are executed: synchronous tasks block execution, while asynchronous tasks allow other operations to run in the meantime.",
		"In object-oriented programming, encapsulation is a principle where the implementation details of a class are hidden, exposing only necessary interfaces for the user.",
		"The Model-View-Controller (MVC) architecture separates an application into three components: the Model manages data, the View displays the user interface, and the Controller handles input logic.",
		"Data structures like arrays, linked lists, and hash tables are fundamental for organizing and storing data efficiently, enabling faster algorithms.",
		"RESTful APIs follow the principles of stateless communication and resource-based architecture, commonly using HTTP methods like GET, POST, PUT, and DELETE.",
		"Machine learning involves training algorithms on data to make predictions or decisions without being explicitly programmed. Key techniques include supervised learning, unsupervised learning, and reinforcement learning.",
		"In cloud computing, scalability is the ability to dynamically adjust resources to handle varying workloads efficiently.",
		"Version control systems like Git enable teams to collaborate on codebases, maintain history, and manage changes through branching and merging.",
		"Cybersecurity practices include encryption, two-factor authentication, and firewalls to protect data and systems from unauthorized access.",
		"The Internet of Things (IoT) connects devices via the internet, enabling them to collect and exchange data for smarter automation and analytics.",
	]
	
	texts = random.choice([motivational_texts, tech_texts])
	return "\n\n".join(random.choice(texts) for _ in range(paragraphs))


def create_meaningful_or_tech_docx(file_path, paragraphs=3):
	"""
	Create a .docx file with meaningful or tech-related text.
	"""
	doc = Document()
	heading = random.choice(["Inspiring Thoughts and Lessons", "Technical Insights"])
	doc.add_heading(heading, level=1)
	doc.add_paragraph(generate_meaningful_or_tech_text(paragraphs))
	doc.save(file_path)


def generate_random_hierarchy_with_tech_and_meaningful_text(base_dir, num_folders, max_num_files, max_depth=3):
	"""
	Generate random folder hierarchy and .docx files with tech and meaningful text.
	"""
	for _ in range(num_folders):
		# Create random folder path
		folder_path = base_dir
		folder_depth = random.randint(1, max_depth)
		for _ in range(folder_depth):
			folder_name = random.choice([
				"Bank", "Other", "General Education", "English", "Personal",
				"Dev", "Job Searching", "Games", "Orders", "tmp",
				"Books", "Portfolio", "Translations", "Treadmill", "Certificates",
				"Personal Development", "Algorithms", "Practical Solutions", "Big Ideas", "Quantum Computing"
			])
			folder_path = os.path.join(folder_path, folder_name)
		os.makedirs(folder_path, exist_ok=True)
		
		# Create random .docx files
		num_files = random.randint(1, max_num_files)
		for file_index in range(num_files):
			file_name = f"{random.choice(['Inspirational Thoughts', 'Tech Guide', 'Learning Notes', 'Code Explained', 'Growth Ideas', 'AI Research', 'Programming Insights', 'Knowledge Base', 'Life Lessons', 'Techniques Overview', 'Practical Tips', 'Philosophy Quotes', 'Data Strategies', 'Innovation Brief', 'Future Plans', 'Creative Solutions', 'Advanced Concepts', 'Simple Hacks', 'Productivity Boosters', 'Visionary Ideas'])}.docx"
			file_path = os.path.join(folder_path, file_name)
			create_meaningful_or_tech_docx(file_path, paragraphs=random.randint(3, 5))
			print(f"    - Created file: {file_path}")


if __name__ == "__main__":
	# Parse command-line arguments
	parser = argparse.ArgumentParser(description="Generate random folder hierarchy with meaningful and technical .docx files.")
	parser.add_argument(
		"base_dir",
		type=str,
		help="Base directory where the folders and files will be generated."
	)
	parser.add_argument(
		"--num_folders", 
		type=int, 
		default=10, 
		help="Number of folders to generate (default: 10)."
	)
	parser.add_argument(
		"--max_num_files", 
		type=int, 
		default=5, 
		help="Maximum number of .docx files per folder (default: 5)."
	)
	parser.add_argument(
		"--max_depth", 
		type=int, 
		default=3, 
		help="Maximum depth of the folder hierarchy (default: 3)."
	)

	args = parser.parse_args()
	# Use the arguments in the script
	base_dir = Path(system_path(args.base_dir))
	print("base_dir: ", base_dir)
	if Path(base_dir).exists():
		raise FileExistsError(f"The directory '{args.base_dir}' already exists. Please provide a non-existing directory.")
	Path(base_dir).mkdir(exist_ok=True)
	generate_random_hierarchy_with_tech_and_meaningful_text(
		base_dir=base_dir,
		num_folders=args.num_folders,
		max_num_files=args.max_num_files,
		max_depth=args.max_depth
	)
	print(f"Random folder hierarchy with tech and meaningful .docx files generated in '{args.base_dir}'.")
